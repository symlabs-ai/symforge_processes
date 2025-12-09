#!/usr/bin/env python3
"""
Exporta Markdown (.md) para DOCX.

Requer a biblioteca opcional:
- python-docx

Uso:
  python symbiotas/mdd_publisher/scripts/export_docx.py \
         --input project/docs/sumario_executivo.md \
         [--output project/output/docs/sumario_executivo.docx]
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path

SCRIPT_DIR = Path(__file__).parent
UTILS_DIR = SCRIPT_DIR / "utils"
if str(UTILS_DIR) not in sys.path:
    sys.path.insert(0, str(UTILS_DIR))

from helpers import (
    MissingDependencyError,
    default_output_for_md,
    log_export,
    md_to_html_basic,
    read_text,
)


def export_docx(input_md: Path, output_docx: Path | None = None) -> Path:
    """
    Exporta Markdown para DOCX com formatação preservada.

    Converte MD para HTML intermediário e então extrai estrutura para DOCX,
    preservando títulos, listas, negrito e itálico.

    Args:
        input_md: Caminho do arquivo .md de entrada
        output_docx: Caminho opcional do .docx de saída

    Returns:
        Path do arquivo DOCX gerado

    Raises:
        SystemExit: Se python-docx não estiver disponível
        FileNotFoundError: Se arquivo de entrada não existir
    """
    try:
        from docx import Document  # type: ignore
    except Exception as e:  # pragma: no cover
        raise MissingDependencyError("Biblioteca 'python-docx' não disponível.") from e

    # Tenta usar BeautifulSoup para melhor parsing
    try:
        from bs4 import BeautifulSoup  # type: ignore
        use_html_parser = True
    except ImportError:
        use_html_parser = False

    text = read_text(input_md)
    out_root = Path("project/output/docs")
    out_path = output_docx or default_output_for_md(input_md, out_root, ".docx")

    doc = Document()

    if use_html_parser:
        # Método avançado: converte MD -> HTML -> DOCX
        html = md_to_html_basic(text)
        soup = BeautifulSoup(html, 'html.parser')

        for elem in soup.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'ul', 'ol', 'li', 'blockquote', 'hr']):
            if elem.name == 'h1':
                doc.add_heading(elem.get_text(strip=True), level=1)
            elif elem.name == 'h2':
                doc.add_heading(elem.get_text(strip=True), level=2)
            elif elem.name == 'h3':
                doc.add_heading(elem.get_text(strip=True), level=3)
            elif elem.name == 'h4':
                doc.add_heading(elem.get_text(strip=True), level=4)
            elif elem.name == 'p' and elem.parent.name not in ['li', 'blockquote']:
                _add_formatted_paragraph(doc, elem)
            elif elem.name == 'blockquote':
                para = doc.add_paragraph(elem.get_text(strip=True))
                para.style = 'Intense Quote'
            elif elem.name in ['ul', 'ol']:
                for li in elem.find_all('li', recursive=False):
                    doc.add_paragraph(li.get_text(strip=True), style='List Bullet' if elem.name == 'ul' else 'List Number')
            elif elem.name == 'hr':
                doc.add_paragraph('─' * 50)
    else:
        # Fallback simples: converte linha a linha (modo legado melhorado)
        import re
        for line in text.splitlines():
            line = line.rstrip()
            if not line:
                continue
            # Títulos
            if line.startswith('#### '):
                doc.add_heading(line[5:].strip(), level=4)
            elif line.startswith('### '):
                doc.add_heading(line[4:].strip(), level=3)
            elif line.startswith('## '):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:].strip(), level=1)
            # Listas
            elif re.match(r'^\s*[\-\*]\s+', line):
                doc.add_paragraph(re.sub(r'^\s*[\-\*]\s+', '', line), style='List Bullet')
            elif re.match(r'^\s*\d+\.\s+', line):
                doc.add_paragraph(re.sub(r'^\s*\d+\.\s+', '', line), style='List Number')
            # Linha horizontal
            elif line.strip() == '---':
                doc.add_paragraph('─' * 50)
            # Parágrafo normal
            else:
                doc.add_paragraph(line)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    log_export(f"DOCX exportado: {input_md} -> {out_path}")
    return out_path


def _add_formatted_paragraph(doc, elem):
    """Adiciona parágrafo com formatação inline (bold, italic)."""
    para = doc.add_paragraph()
    for content in elem.children:
        if hasattr(content, 'name'):
            run = para.add_run(content.get_text())
            if content.name == 'strong':
                run.bold = True
            elif content.name == 'em':
                run.italic = True
            elif content.name == 'code':
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
        else:
            para.add_run(str(content))


def main() -> int:
    ap = argparse.ArgumentParser(description="MDD Publisher - Exportar Markdown para DOCX")
    ap.add_argument("--input", required=True, help="Caminho do arquivo .md de entrada")
    ap.add_argument("--output", required=False, help="Caminho do .docx de saída")
    args = ap.parse_args()

    in_path = Path(args.input)
    if not in_path.exists():
        print(f"[ERRO] Arquivo de entrada não encontrado: {in_path}", file=sys.stderr)
        return 2
    out_path = Path(args.output) if args.output else None
    try:
        final_path = export_docx(in_path, out_path)
        print(str(final_path))
        return 0
    except MissingDependencyError as me:
        print(f"[ERRO] {me}", file=sys.stderr)
        return 1
    except Exception as exc:
        log_export(f"FALHA ao exportar DOCX: {in_path} - {exc}")
        print(f"[ERRO] Falha ao exportar DOCX: {exc}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
